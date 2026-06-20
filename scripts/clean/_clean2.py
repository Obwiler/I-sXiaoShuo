import os
import re
from docx import Document
path = os.path.join(PROJECT, "设定", "02_大纲章纲.docx")
doc = Document(path)
c=0
for p in doc.paragraphs:
    if "——" in p.text:
        c+=1
        text = p.text.replace("——","。")
        text = re.sub(r"。。+","。",text)
        for run in p.runs: run.text=""
        if p.runs: p.runs[0].text=text
        else: p.add_run(text)
doc.save(path)
d2=Document(path)
r1=sum(1 for p in d2.paragraphs if "——" in p.text)
r2=sum(1 for p in d2.paragraphs if re.search(r"不是[^。]{2,50}(而是|只是)",p.text))
print(f"Cleaned {c} paragraphs. Remaining: dash={r1} not-but={r2}")
