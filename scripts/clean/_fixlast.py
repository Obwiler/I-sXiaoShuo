import os
import re
from docx import Document
path = os.path.join(PROJECT, "设定", "02_大纲章纲.docx")
doc = Document(path)
for i, p in enumerate(doc.paragraphs):
    m = re.search(r"不是[^。]{2,50}(而是|只是)", p.text)
    if m:
        print(f"[{i}] ...{p.text[max(0,m.start()-20):m.end()+20]}...")
        # Fix it
        text = p.text
        # Remove the negation framing
        text = re.sub(r"不是[^。]{2,50}(而是|只是)[^。]{2,50}", lambda x: x.group(0).split("而是" if "而是" in x.group(0) else "只是")[-1] if "而是" in x.group(0) or "只是" in x.group(0) else x.group(0), text)
        for run in p.runs: run.text=""
        if p.runs: p.runs[0].text=text
        else: p.add_run(text)
doc.save(path)
print("Fixed")
