import os
import docx
from docx.shared import Pt, Mm
from docx.enum.text import WD_LINE_SPACING,  WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

files = [
    "00_设定总纲.docx", "01_人物档案.docx", "02_大纲章纲.docx",
    "03_伏笔追踪.docx", "04_世界观与势力.docx", "第1卷_时间线.docx",
]

def sf(run, fn, sz, b):
    run.font.name = fn; run.font.size = sz; run.font.bold = b
    rp = run._element.get_or_add_rPr()
    rf = rp.find(qn("w:rFonts"))
    if rf is None:
        rf = docx.oxml.OxmlElement("w:rFonts")
        rp.append(rf)
    rf.set(qn("w:eastAsia"), fn)

def is_h1(t):
    if len(t)>30: return False
    return any(re.match(p,t) for p in [
        r"^第[一二三四五六七八九十\d]+卷", r"^\u7b2c\d+\u7ae0\u300c",
        r"^\u7b2c\d+\u7ae0\s+\u589e\u8865\u8282",
        r"^第[一二三四五六七八九十\d]+章\s",
        r"^\u5f27\u6bb5[一二三四五六七八九十\d]+",
        r"^附录", r"^\u5e8f\u7ae0$", r"^\u7ec8\u7ae0$",
        r"^\u300a\u5965\u8d1d\u7ef4\u52d2\u300b", r"^\u5168\u5377"])

def is_h2(t):
    if len(t)>40: return False
    return any(re.match(p,t) for p in [
        r"^节[一二三四五六七八九十\d]+", r"^[一二三四五六七八九十]+[、\.\s]",
        r"^第[一二三四五六七八九十\d]+章", r"^正文时间锚定",
        r"^说明", r"^世界版图", r"^七大洲", r"^定义", r"^参考公历",
        r"^文历元年", r"^辰夏在文历", r"^国名源流考",
        r"^事件发展树", r"^主线事件树", r"^人物成长树",
        r"^国际大局事件树", r"^压制者技术树", r"^时间线使用规则",
        r"^纪元对照表", r"^平行纪年对照"])

def is_h3(t):
    if len(t)>50: return False
    return any(re.match(p,t) for p in [
        r"^\uff08[一二三四五六七八九十\d]+\uff09",
        r"^核心机构", r"^与其他势力", r"^内部教义分化",
        r"^压制装置技术来源"])

def fmt_tbl(tbl):
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    nrows = len(tbl.rows)
    for ri, row in enumerate(tbl.rows):
        for cell in row.cells:
            tc = cell._tc; tp = tc.get_or_add_tcPr()
            tb = tp.find(qn("w:tcBorders"))
            if tb is not None: tp.remove(tb)
            nb = docx.oxml.OxmlElement("w:tcBorders")
            if ri == 0:
                nb.append(parse_xml('<w:top ' + nsdecls("w") +
                    ' w:val="single" w:sz="12" w:space="0" w:color="000000"/>'))
                nb.append(parse_xml('<w:bottom ' + nsdecls("w") +
                    ' w:val="single" w:sz="4" w:space="0" w:color="000000"/>'))
            if ri == nrows - 1 and ri != 0:
                nb.append(parse_xml('<w:bottom ' + nsdecls("w") +
                    ' w:val="single" w:sz="12" w:space="0" w:color="000000"/>'))
            tp.append(nb)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                para.paragraph_format.line_spacing = Pt(15.6)
                for run in para.runs:
                    sf(run, "黑体" if ri==0 else "方正仿宋_GB2312",
                       Pt(14) if ri==0 else Pt(10.5), ri==0)

for fname in files:
    path = os.path.join(PROJECT, "设定", fname)
    doc = docx.Document(path)
    for sec in doc.sections:
        sec.top_margin = Mm(37); sec.bottom_margin = Mm(35)
        sec.left_margin = Mm(28); sec.right_margin = Mm(26)
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        memo = t.startswith("记忆点") or t.startswith("备注")
        h1 = is_h1(t)
        h2 = is_h2(t) and not h1
        h3 = is_h3(t) and not h1 and not h2
        if h1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            for r in p.runs: sf(r, "黑体", Pt(16), True)
        elif h2:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            for r in p.runs: sf(r, "黑体", Pt(14), True)
        elif h3:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            for r in p.runs: sf(r, "楷体", Pt(12), False)
        elif memo:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(0)
            for r in p.runs: sf(r, "方正仿宋_GB2312", Pt(9), False)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(21)
            for r in p.runs: sf(r, "方正仿宋_GB2312", Pt(10.5), False)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(15.6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    for tbl in doc.tables:
        fmt_tbl(tbl)
    doc.save(path)
    print(fname + ": " + str(len(doc.paragraphs)) + "p " + str(len(doc.tables)) + "t")
print("ALL DONE")
