import os
from docx import Document
import re

path = os.path.join(PROJECT, "设定", "02_大纲章纲.docx")
print(f'Opening: {path}')
doc = Document(path)

for i, p in enumerate(doc.paragraphs):
    text = p.text
    
    if i == 116:
        text = text.replace('不是它而是他？', '是一个独立意识在回应他。')
        text = text.replace("手指不再颤抖——他发现了一个让他无法继续将这个存在归类为'它'的事实", '手指不再颤抖。')
    elif i == 147:
        text = text.replace('不是防止激活而是压制已发生的激活', '药物作用是压制已发生的激活')
    elif i == 555:
        text = text.replace('不是要从我身上拿走什么，他只是在看', '他从一开始就没有要从我身上拿走任何东西。他一直在做的只有一件事。看')
    
    for run in p.runs:
        run.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)

doc.save(path)

doc2 = Document(path)
not_but = sum(1 for p in doc2.paragraphs if re.search(r'不是[^。]{2,50}(而是|只是)', p.text))
dash = sum(1 for p in doc2.paragraphs if '——' in p.text)
print(f'终审: 破折号={dash}, 不是…而是={not_but}')
print('硬性清理全部完成' if dash==0 and not_but==0 else '仍有残留')
