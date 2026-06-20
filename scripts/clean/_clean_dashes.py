import os
import re
from docx import Document

path = os.path.join(PROJECT, "设定", "02_大纲章纲.docx")
doc = Document(path)

count = 0
for p in doc.paragraphs:
    if "——" in p.text:
        count += 1
        # Replace all em dashes with period
        text = p.text.replace("——", "。")
        # Clean up doubled periods
        text = re.sub(r"。。+", "。", text)
        text = re.sub(r"。，", "，", text)
        text = re.sub(r"，。", "。", text)
        # Clean period at start
        text = re.sub(r"^。", "", text)
        # Update
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)

doc.save(path)
print(f"Cleaned {count} paragraphs with em dashes")

# Verify
doc2 = Document(path)
remaining = sum(1 for p in doc2.paragraphs if "——" in p.text)
notbut = sum(1 for p in doc2.paragraphs if re.search(r"不是[^。]{2,50}(而是|只是)", p.text))
print(f"Remaining em dashes: {remaining}")
print(f"Remaining not-but: {notbut}")
print("All clean!" if remaining == 0 and notbut == 0 else "Still needs work")
