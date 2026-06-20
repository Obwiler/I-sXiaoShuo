
import docx, re, os

PROJECT = os.environ.get("WEBNOVEL_PROJECT") or os.getcwd()
SETTINGS_DIR = os.path.join(PROJECT, "设定")
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

files = [f for f in os.listdir(SETTINGS_DIR) if f.endswith('.docx') and os.path.isfile(os.path.join(SETTINGS_DIR, f))]
if not files:
    print(f'警告: 设定目录下未找到 .docx 文件: {SETTINGS_DIR}')
    sys.exit(1)

body_font = '仿宋'  # 仿宋
body_size = Pt(16)
title_font = '黑体'  # 黑体
title_size = Pt(18)
line_sp = Pt(28)
indent_val = Pt(32)

def is_title_line(text):
    if len(text) > 30:
        return False
    p1 = r'^第[一二三四五六七八九十\d]+卷'
    p2 = r'^第\d+章「|『|“'  # 第XXX章「
    p3 = r'^第\d+章\s+增补节'
    p4 = r'^节[一二三四五六七八九十\d]+'
    p5 = r'^弧段[一二三四五六七八九十\d]+'
    p6 = r'^附录'
    p7 = r'^《.*?》'  # 任意项目名称
    p8 = r'^第[一二三四五六七八九十\d]+章\s'
    p9 = r'^[一二三四五六七八九十]+[、\.\s]'  # 一、二、三 等
    patterns = [p1, p2, p3, p4, p5, p6, p7, p8, p9]
    for pat in patterns:
        if re.match(pat, text):
            return True
    return False

for fname in files:
    path = os.path.join(PROJECT, "设定", fname)
    doc = docx.Document(path)
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        is_memo = text.startswith('记忆点')  # 记忆点
        is_title = is_title_line(text)
        if is_title:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            for r in p.runs:
                r.font.name = title_font
                r.font.size = title_size
                r.font.bold = True
                rPr2 = r._element.get_or_add_rPr()
                rFonts2 = rPr2.find(qn('w:rFonts'))
                if rFonts2 is None:
                    rFonts2 = docx.oxml.OxmlElement('w:rFonts')
                    rPr2.append(rFonts2)
                rFonts2.set(qn('w:eastAsia'), title_font)
        elif is_memo:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(0)
            for r in p.runs:
                r.font.name = body_font
                r.font.size = Pt(14)
                r.font.bold = False
                rPr2 = r._element.get_or_add_rPr()
                rFonts2 = rPr2.find(qn('w:rFonts'))
                if rFonts2 is None:
                    rFonts2 = docx.oxml.OxmlElement('w:rFonts')
                    rPr2.append(rFonts2)
                rFonts2.set(qn('w:eastAsia'), body_font)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = indent_val
            for r in p.runs:
                r.font.name = body_font
                r.font.size = body_size
                r.font.bold = False
                rPr2 = r._element.get_or_add_rPr()
                rFonts2 = rPr2.find(qn('w:rFonts'))
                if rFonts2 is None:
                    rFonts2 = docx.oxml.OxmlElement('w:rFonts')
                    rPr2.append(rFonts2)
                rFonts2.set(qn('w:eastAsia'), body_font)
        p.paragraph_format.line_spacing = line_sp
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
    doc.save(path)
    print(fname + ': done')

# Verify
for fname in files:
    path = os.path.join(PROJECT, "设定", fname)
    doc = docx.Document(path)
    bad = 0
    for p in doc.paragraphs[:40]:
        t = p.text.strip()
        if not t or len(t) < 3:
            continue
        r = p.runs[0] if p.runs else None
        if r and r.font.name == title_font and len(t) > 25 and '本文档' not in t:
            bad += 1
            if bad <= 3:
                print('  OVER-TAGGED: ' + t[:50])
    if bad == 0:
        print('  [' + fname + '] title detection correct')
