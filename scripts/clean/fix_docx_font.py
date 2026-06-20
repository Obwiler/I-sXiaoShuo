import os
import docx
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

files = [
    "00_设定总纲.docx", "01_人物档案.docx", "02_大纲章纲.docx",
    "03_伏笔追踪.docx", "04_世界观与势力.docx", "第1卷_时间线.docx",
]

BF = "方正仿宋_GB2312"

def sf(run, fn, sz, b):
    run.font.name = fn; run.font.size = sz; run.font.bold = b
    rp = run._element.get_or_add_rPr()
    rf = rp.find(qn("w:rFonts"))
    if rf is None:
        rf = docx.oxml.OxmlElement("w:rFonts")
        rp.append(rf)
    rf.set(qn("w:eastAsia"), fn)

for fname in files:
    path = os.path.join(PROJECT, "设定", fname)
    doc = docx.Document(path)
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        for r in p.runs:
            if r.font.name and r.font.name in ("方正仿宋", "仿宋"):
                sf(r, BF, r.font.size, r.font.bold)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for r in para.runs:
                        if r.font.name and r.font.name in ("方正仿宋", "仿宋"):
                            sf(r, BF, r.font.size, r.font.bold)
    style = doc.styles["Normal"]
    style.font.name = BF
    rPr = style.element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = docx.oxml.OxmlElement("w:rFonts")
        rPr.append(rf)
    rf.set(qn("w:eastAsia"), BF)
    
    doc.save(path)
    print(fname + ": " + str(len(doc.paragraphs)) + "p")
print("ALL DONE")
