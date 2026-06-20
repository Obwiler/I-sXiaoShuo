import os
from docx import Document
import re

path = os.path.join(PROJECT, "设定", "02_大纲章纲.docx")
doc = Document(path)

dash_p = 0
notbut_p = 0
total_p = len(doc.paragraphs)

for p in doc.paragraphs:
    has_change = False
    text = p.text
    
    if '——' in text:
        text = text.replace('——', '。')
        has_change = True
        dash_p += 1
    
    # Check for not-but
    if re.search(r'不是[^。]{2,50}(而是|只是)', text):
        text = re.sub(r'不是[^。]{2,50}(而是|只是)[^。]{2,80}', lambda m: m.group(0).split('而是' if '而是' in m.group(0) else '只是')[-1] if any(k in m.group(0) for k in ['而是','只是']) else m.group(0), text)
        has_change = True
        notbut_p += 1
    
    if has_change:
        text = re.sub(r'。。+', '。', text)
        text = re.sub(r'^。', '', text)
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)

doc.save(path)

# Verify
doc2 = Document(path)
dash_r = sum(1 for p in doc2.paragraphs if '——' in p.text)
nr_r = sum(1 for p in doc2.paragraphs if re.search(r'不是[^。]{2,50}(而是|只是)', p.text))

print(f'总段落: {total_p}')
print(f'破折号清理: {dash_p} -> 剩余 {dash_r}')
print(f'不是…而是清理: {notbut_p} -> 剩余 {nr_r}')

# Count chapters and sections
ch_count = 0
sec_count = 0
for p in doc2.paragraphs:
    t = p.text.strip()
    if '第' in t and '章' in t and ('★★' in t or '★★★' in t or '★★★★' in t):
        ch_count += 1
    if t.startswith(('节一','节二','节三','节四','节五')):
        sec_count += 1

print(f'章节统计: {ch_count} 章, {sec_count} 节 (约{ch_count*4}节平均)')
print('清理完成' if dash_r == 0 and nr_r == 0 else '仍有残留')
