import re, os

PROJECT = os.environ.get("WEBNOVEL_PROJECT") or os.getcwd()
from docx import Document

# Only check the outline
fp = os.path.join(PROJECT, "设定", "02_大纲章纲.docx")
doc = Document(fp)

total = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    if '不是' in text and ('而是' in text or '只是' in text):
        idx = text.find('不是')
        rest = text[idx:idx+120]
        if '而是' in rest or '只是' in rest:
            print(f'大纲 L{i+1}: ...{rest[:100]}...')
            total += 1

if total == 0:
    print('✅ 大纲零残留')
else:
    print(f'\n⚠️ 大纲发现 {total} 处')

# Now check standalone "不是XX" patterns in outline
print('\n--- 检查大纲中的"不是XX"独立成句 ---')
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    lines = text.split('。')
    for j, sent in enumerate(lines):
        s = sent.strip()
        if s.startswith('不是') and len(s) < 40:
            print(f'大纲 L{i+1} 句{j+1}: "{s}"')
            total += 1

if total == 0:
    print('✅ 大纲零残留（含独立成句）')
else:
    print(f'\n⚠️ 总计 {total} 处')
