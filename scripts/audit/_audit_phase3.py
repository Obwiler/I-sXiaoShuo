from docx import Document

# Check 00_设定总纲.docx for new chapter
doc = Document(os.path.join(PROJECT, "设定", "00_设定总纲.docx"))
print('=== 00_设定总纲.docx 段落索引 ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        if len(t) < 80:
            print(f'P{i}: {t}')
        else:
            print(f'P{i}: {t[:80]}...')
pr
print()

# Check 04_世界观与势力.docx for new appendix
doc2 = Document(os.path.join(PROJECT, "设定", "04_世界观与势力.docx"))
print('=== 04_世界观与势力.docx 段落索引 ===')
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    if t:
        if len(t) < 100:
            print(f'P{i}: {t}')
        else:
            print(f'P{i}: {t[:100]}...')
